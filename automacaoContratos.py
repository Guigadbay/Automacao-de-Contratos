'''
Passos:
- Passar as informações da planilha para o arquivo word.
- Salvar essa arquivo em uma pasta espeífica (contratos).
- Repetir para todas a planilha.

Bibliotecas necessárias:
Openpyxl: Ler planilhas.
Python-docx: Cria arquivos Word.
'''
from openpyxl import load_workbook #Permite abrir planilhas.
from docx import Document #Permite abrir e criar docx.
from datetime import datetime #Permite inserir a data atual no contrato.

planilha_fornecedores = load_workbook('./fornecedores.xlsx') #Essa váriavel serve para abrir a planilha fornecedores. 
pagina_fornecedores = planilha_fornecedores['Sheet1'] #Especificando que é a página sheet1 do arquivo fornecedores.

#iter_rows função do Openpyxl que permite ler cada linha da planilha seguindo os parametros dos parentesis.
#min_row diz que é para começar a ler no minimo a partir da linha 2.
#values_only=True diz que é para ler somente valres verdadeiros, ou seja, o que for texto.
for linha in pagina_fornecedores.iter_rows(min_row=2,values_only=True):
    nome_empresa, endereco, cidade, estado, cep, telefone, email, setor = linha #Extraindo os dados de cada linha da planilha e atribuindo a uma variavel.

    arquivo_word = Document() #Criando e defenindo o arquivo word a variavel arquivo_word.
    arquivo_word.add_heading('Contrato de Prestação de Serviço',0) #Comando do Pythondox para escrever um cabeçalho e o 0 é uma forma predefinida de estilização para ele .

    #Criando o texto do contrato e substituindo as informaçãoes que sempre irão mudar pelas variáveis que eu criei acima.
    texto_contrato = f""" 
    Este contrato de prestação de serviços é feito entre {nome_empresa}, com endereço em {endereco}, 
    {cidade}, {estado}, CEP {cep}, doravante denominado FORNECEDOR, e a empresa CONTRATANTE.

    Pelo presente instrumento particular, as partes têm, entre si, justo e acordado o seguinte:

    1. OBJETO DO CONTRATO
    O FORNECEDOR compromete-se a fornecer à CONTRATANTE os serviços/material de acordo com as especificações acordadas, respeitando os padrões de qualidade e os prazos estipulados.

    2. PRAZO
    Este contrato tem prazo de vigência de 12 (doze) meses, iniciando-se na data de sua assinatura, podendo ser renovado conforme acordo entre as partes.

    3. VALOR E FORMA DE PAGAMENTO
    O valor dos serviços prestados será acordado conforme as demandas da CONTRATANTE e a capacidade de entrega do FORNECEDOR. Os pagamentos serão realizados mensalmente, mediante apresentação de nota fiscal.

    4. CONFIDENCIALIDADE
    Todas as informações trocadas entre as partes durante a vigência deste contrato serão tratadas como confidenciais.

    Para firmeza e como prova de assim haverem justo e contratado, as partes assinam o presente contrato em duas vias de igual teor e forma.

    FORNECEDOR: {nome_empresa}
    E-mail: {email}

    CONTRATANTE: [NOME CONTRATANTE]
    E-mail: [E-MAIL CONTRATANTE]

    [CIDADE], {datetime.now().strftime('%d/%m/%Y')} 
    """
    #O comando acima diz que é para escrever no contrato a data de hoje no formato brasileiro de dia, mes e ano.
    arquivo_word.add_paragraph(texto_contrato) #Comando do Pythondox para escrever o parágrafo acima.

    arquivo_word.save(f'./contratos/contrato_{nome_empresa}.docx') #Comando do Pythondox para salvar esse arquivo word na pasta contratos, com o nome de contratos_nome da empresa.docx.
